#!/usr/bin/env python
from docx import Document


class GenericTemplate(object):

    def __init__(self):
        self.document = Document('../../assets/word-base/dissertate.docx')

    def fill(self):

        # Cover page
        self.document.add_heading('Title of the dissertation', level=1)

        list = ['a dissertation presented',
                'by',
                'Firstname M. Lastname',
                'to',
                'the department of motor vehicles',
                '',
                'in partial fulfillment of the requirements',
                'for the degree of',
                'Doctor of Philosophy',
                'in the subject of',
                'Psychology',
                '',
                'Harvard University',
                'Cambridge, Massachusetts',
                'May 2014']

        for i in list:
            self.document.add_paragraph(i, style="CoverBody")

        self.document.add_page_break()

        p = self.document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        self.document.add_heading('Heading, level 1', level=1)
        self.document.add_paragraph('Intense quote', style='IntenseQuote')

        self.document.add_paragraph(
            'first item in unordered list', style='ListBullet'
        )
        self.document.add_paragraph(
            'first item in ordered list', style='ListNumber'
        )

        #document.add_picture('monty-truth.png', width=Inches(1.25))

    def save(self):
        self.document.save('demo.docx')
